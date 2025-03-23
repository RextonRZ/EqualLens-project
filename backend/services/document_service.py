import os
import logging
from typing import Dict, Tuple, Any
from io import BytesIO
import base64
from dotenv import load_dotenv
from google.api_core.client_options import ClientOptions
from google.cloud import documentai

# Load environment variables
load_dotenv()

# Configure logging
logger = logging.getLogger(__name__)

# Check for optional document conversion libraries
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    reportlab_available = True
except ImportError:
    reportlab_available = False
    logger.warning("reportlab not installed. Install with: pip install reportlab")

try:
    import docx
    docx_available = True
except ImportError:
    docx_available = False
    logger.warning("python-docx not installed. Install with: pip install python-docx")


class DocumentService:
    """Service for processing documents and extracting data."""
    
    @staticmethod
    def convert_to_pdf(file_content: bytes, file_name: str) -> tuple:
        """Extracts text from DOC/DOCX files and creates a PDF with the extracted text.
        
        Args:
            file_content: The file content as bytes
            file_name: Original filename with extension
            
        Returns:
            Tuple of (converted PDF content as bytes, mime_type)
        """
        file_extension = os.path.splitext(file_name)[1].lower()
        
        if file_extension not in ['.doc', '.docx']:
            # No conversion needed, return original content
            return file_content, "application/pdf"
        
        # Check if required libraries are available
        if not reportlab_available:
            logger.error("reportlab not installed. Cannot convert document to PDF.")
            return file_content, f"application/{file_extension.replace('.', '')}"
        
        # Create BytesIO for the file content
        file_bytesio = BytesIO(file_content)
        extracted_text = ""
        
        try:
            # Process DOCX files
            if file_extension == '.docx' and docx_available:
                logger.info(f"Converting DOCX file: {file_name}")
                document = docx.Document(file_bytesio)
                
                # Extract text from paragraphs
                paragraphs = []
                for para in document.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                
                # Extract text from tables
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if para.text.strip():
                                    paragraphs.append(para.text)
                
                extracted_text = "\n".join(paragraphs)
                logger.info(f"Successfully extracted {len(paragraphs)} paragraphs from DOCX")
                
            # Process DOC files (limited support)
            elif file_extension == '.doc':
                extracted_text = "NOTE: This is a DOC file that has been converted to plain text.\n\n"
                
                # Try to extract some text with basic encoding detection
                try:
                    import chardet
                    encoding_result = chardet.detect(file_content)
                    encoding = encoding_result['encoding'] if encoding_result['confidence'] > 0.5 else 'utf-8'
                    text = file_content.decode(encoding, errors='ignore')
                    extracted_text += text
                    logger.info(f"Extracted text from DOC file using encoding {encoding}")
                except Exception as e:
                    logger.warning(f"Could not extract text from DOC file: {e}")
                    extracted_text += "Could not extract text content from this DOC file."
            
            # Create PDF with extracted text
            pdf_bytesio = BytesIO()
            pdf = SimpleDocTemplate(pdf_bytesio, pagesize=letter)
            styles = getSampleStyleSheet()
            
            # Build content for PDF
            content = []
            content.append(Paragraph(f"Extracted Text from: {file_name}", styles['Title']))
            content.append(Spacer(1, 20))
            
            # Split the text into manageable paragraphs
            paragraphs = extracted_text.split('\n')
            for i, para_text in enumerate(paragraphs):
                if i > 1000:  # Limit number of paragraphs for memory safety
                    content.append(Paragraph("... (text truncated due to size) ...", styles['Normal']))
                    break
                    
                if para_text.strip():
                    try:
                        # Clean the text to prevent XML parsing errors in reportlab
                        cleaned_text = para_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        p = Paragraph(cleaned_text, styles['Normal'])
                        content.append(p)
                        content.append(Spacer(1, 10))
                    except Exception as e:
                        logger.warning(f"Could not add paragraph to PDF: {e}")
            
            # Build PDF
            pdf.build(content)
            
            # Get the PDF content
            pdf_content = pdf_bytesio.getvalue()
            pdf_bytesio.close()
            
            logger.info(f"Successfully created PDF with {len(content)} content elements, size: {len(pdf_content)} bytes")
            return pdf_content, "application/pdf"
            
        except Exception as e:
            logger.error(f"Error converting document to PDF: {str(e)}")
            # Return original content if conversion fails
            return file_content, f"application/{file_extension.replace('.', '')}"
    
    @staticmethod
    def process_document(file_content: bytes, mime_type: str, file_name: str) -> Dict[str, Any]:
        """Processes a document using an existing Document AI processor and extracts structured data."""
        project_id = os.getenv("DOCUMENTAI_PROJECT_ID", "default_project_id")
        location = os.getenv("DOCUMENTAI_LOCATION", "us")  # "us" or "eu"
        processor_id = os.getenv("DOCUMENTAI_PROCESSOR_ID", "default_processor_id")
        processor_version = os.getenv("DOCUMENTAI_PROCESSOR_VERSION", "default_processor_version")
        
        # Get file extension
        file_extension = os.path.splitext(file_name)[1].lower()
        
        # Convert doc/docx to PDF if needed
        if file_extension in ['.doc', '.docx']:
            logger.info(f"Converting {file_extension} file to PDF for Document AI processing")
            file_content, converted_mime_type = DocumentService.convert_to_pdf(file_content, file_name)
            mime_type = converted_mime_type  # Make sure to use the new mime type
            logger.info(f"Conversion complete. New MIME type: {mime_type}, content size: {len(file_content)} bytes")

        # Define API endpoint
        opts = ClientOptions(api_endpoint=f"{location}-documentai.googleapis.com")

        # Log Document AI request details
        logger.info(f"Sending document to Document AI - Processor ID: {processor_id}, MIME type: {mime_type}")
        
        try:
            # Initialize the Document AI client
            client = documentai.DocumentProcessorServiceClient(client_options=opts)

            # Construct processor resource name with version
            processor_name = f"projects/{project_id}/locations/{location}/processors/{processor_id}/processorVersions/{processor_version}"

            # Create a raw document request
            raw_document = documentai.RawDocument(content=file_content, mime_type=mime_type)

            # Create a request using the existing processor
            request = documentai.ProcessRequest(name=processor_name, raw_document=raw_document)

            # Process the document
            result = client.process_document(request=request)

            # Extract structured data from the document
            document = result.document
            structured_data = {}

            for entity in document.entities:
                field_name = entity.type_  # Field name as defined in the processor
                field_value = entity.mention_text  # Extracted value for the field
                structured_data[field_name] = field_value
                
            logger.info(f"Document AI processing successful. Extracted {len(structured_data)} fields.")
            return structured_data
            
        except Exception as e:
            logger.error(f"Document AI processing failed: {str(e)}")
            # If Document AI fails but we have extracted text from DOCX, provide it as fallback
            if file_extension in ['.doc', '.docx'] and file_content:
                logger.info("Using text extraction as fallback for Document AI")
                return {"extracted_text": "Text extracted during conversion (Document AI processing failed)"}
            raise
